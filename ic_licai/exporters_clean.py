# ic_licai/exporters_clean.py — exports with safe local folder (~/Documents/ICLicAI/reports)

import io, os, json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

try:
    from openpyxl import Workbook
    HAVE_XLSX = True
except Exception:
    HAVE_XLSX = False


def ensure_reports_dir(base: Path = None) -> Path:
    """Ensure ~/Documents/ICLicAI/reports exists."""
    base = Path(base or Path.home() / "Documents" / "ICLicAI" / "reports")
    base.mkdir(parents=True, exist_ok=True)
    return base


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


# --------------------------- DOCX Advisory / Templates ---------------------------
def export_advisory_docx(case: dict, mode: str, reports_dir: Path = None):
    """Create a Word advisory or template docx."""
    reports_dir = ensure_reports_dir(reports_dir)
    fname = {
        "ADVISORY": f"Advisory_Report_{case['case_name']}_{_timestamp()}.docx",
        "TEMPLATE_FRAND": f"Template_FRAND_{case['case_name']}_{_timestamp()}.docx",
        "TEMPLATE_CO_CREATION": f"Template_CoCreation_{case['case_name']}_{_timestamp()}.docx",
        "TEMPLATE_NON_TRADITIONAL": f"Template_NonTraditional_{case['case_name']}_{_timestamp()}.docx",
    }[mode]
    path = reports_dir / fname

    if HAVE_DOCX:
        doc = Document()
        doc.add_heading("Intangible Capital — Advisory / Licensing", level=1)
        doc.add_paragraph(f"Case: {case['case_name']}  |  Size: {case['company_size']}  |  Sector: {case['sector']}")

        if mode == "ADVISORY":
            _docx_advisory_body(doc, case)
        elif mode == "TEMPLATE_FRAND":
            _docx_template_frand(doc)
        elif mode == "TEMPLATE_CO_CREATION":
            _docx_template_cocreation(doc)
        else:
            _docx_template_non_traditional(doc)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        path.write_bytes(bio.getvalue())
        return bio.getvalue(), str(path)

    # Fallback text file
    txt = _plain_advisory_text(case, mode)
    path_txt = str(path).replace(".docx", ".txt")
    Path(path_txt).write_text(txt, encoding="utf-8")
    return txt.encode("utf-8"), path_txt


def _docx_advisory_body(doc, case: dict):
    doc.add_heading("4-Leaf Summary", level=2)
    for k, v in case["four_leaf"].items():
        if v.strip():
            doc.add_paragraph(f"{k.capitalize()}: {v.strip()}")

    doc.add_heading("10-Steps (advisory notes)", level=2)
    for step in case["ten_steps"]:
        if step["notes"].strip():
            doc.add_paragraph(f"{step['label']}: {step['notes'].strip()}")

    doc.add_heading("Licensing Intent & FRAND", level=2)
    lic = case.get("licensing", {})
    if lic.get("intent"):
        doc.add_paragraph(f"Intent: {lic['intent']}")
    if lic.get("frand_notes", "").strip():
        doc.add_paragraph(f"FRAND: {lic['frand_notes'].strip()}")

    if case.get("esg_rows"):
        doc.add_paragraph("ESG CSV present — mapped to IA register (4-Leaf cues).")

    doc.add_heading("IAS 38 Alignment", level=2)
    doc.add_paragraph("Identify, control, and future economic benefit evidenced. Valuation steps held by Areopa (trade-secreted).")


def _docx_template_frand(doc):
    doc.add_heading("Standard FRAND Licence Template (Draft)", level=2)
    doc.add_paragraph("• Grant: Non-exclusive licence for defined field-of-use/territory.")
    doc.add_paragraph("• Fees: Reasonable royalty or fixed fee; MFN across equivalent licensees.")
    doc.add_paragraph("• Audit: Periodic statements; audit rights with notice.")
    doc.add_paragraph("• Termination: Clear triggers; survival of confidentiality.")


def _docx_template_cocreation(doc):
    doc.add_heading("Co-Creation Agreement Template (Draft)", level=2)
    doc.add_paragraph("• Scope: Joint development of artefacts; roles & contributions.")
    doc.add_paragraph("• IP: Background vs Foreground; FRAND access if reused.")
    doc.add_paragraph("• Governance: Steering cadence; change control; acceptance criteria.")
    doc.add_paragraph("• Commercialisation: Licence-back rights; royalty corridors; spin-out options.")


def _docx_template_non_traditional(doc):
    doc.add_heading("Non-Traditional Asset Licence — Codified Knowledge (Draft)", level=2)
    doc.add_paragraph("• Object: Know-how pack (procedures, training, data).")
    doc.add_paragraph("• Terms: Social benefit or commercial use; attribution; open clauses.")
    doc.add_paragraph("• Safeguards: Confidential annex; revocation for misuse; redistribution rules.")
    doc.add_paragraph("• Compatibility: FRAND-style parity; reasonable fees in line with impact.")


def _plain_advisory_text(case: dict, mode: str) -> str:
    head = f"Case: {case['case_name']} | Size: {case['company_size']} | Sector: {case['sector']}\n"
    return head + f"Mode: {mode}"


# ------------------------------- XLSX IA Register -------------------------------
def export_ia_register_xlsx(case: dict, reports_dir: Path = None):
    reports_dir = ensure_reports_dir(reports_dir)
    fname = f"IA_Register_{case['case_name']}_{_timestamp()}.xlsx"
    path = reports_dir / fname

    if HAVE_XLSX:
        wb = Workbook()
        ws = wb.active
        ws.title = "IA Register"
        ws.append(["Case", case["case_name"]])
        ws.append(["Size", case["company_size"]])
        ws.append(["Sector", case["sector"]])
        ws.append([])
        ws.append(["4-Leaf Category", "Notes"])
        for k, v in case["four_leaf"].items():
            if v.strip():
                ws.append([k.capitalize(), v.strip()])
        ws.append([])
        ws.append(["10-Steps", "Notes"])
        for step in case["ten_steps"]:
            if step["notes"].strip():
                ws.append([step["label"], step["notes"].strip()])
        ws.append([])
        ws.append(["Licensing Intent", case.get("licensing", {}).get("intent", "")])
        ws.append(["FRAND Notes", case.get("licensing", {}).get("frand_notes", "")])
        ws.append([])
        ws.append(["ESG CSV present", "Yes" if case.get("esg_rows") else "No"])
        bio = io.BytesIO()
        wb.save(bio)
        bio.seek(0)
        path.write_bytes(bio.getvalue())
        return bio.getvalue(), str(path)

    csv_path = str(path).replace(".xlsx", ".csv")
    lines = [f"Case,{case['case_name']}"]
    Path(csv_path).write_text("\n".join(lines), encoding="utf-8")
    return b"\n".join(lines), csv_path


# -------------------------------- JSON Export --------------------------------
def export_case_json(case: dict, reports_dir: Path = None):
    reports_dir = ensure_reports_dir(reports_dir)
    fname = f"CaseData_{case['case_name']}_{_timestamp()}.json"
    path = reports_dir / fname
    data = json.dumps(case, indent=2, ensure_ascii=False)
    Path(path).write_text(data, encoding="utf-8")
    return data.encode("utf-8"), str(path)
